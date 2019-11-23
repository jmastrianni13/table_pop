##Table populating program
##Version 0.1.6 additions:
##                         Program simplified (no table manipulation allowed moving forward)
##                         openpyxl library replaced with pandas
##                         Enforce order integrity option added
##                         Wider use of functions

import datetime 
import argparse
import os

from docx import Document

import pandas as pd

# TODO: Use now().strftime('%Y%m%d') to get the same result.  Move to inside the __main__ section. 
yyyymmdd = datetime.datetime.now().strftime('%Y%m%d')

def file_loc(folder, file_name, ext):
    """Makes a file path
    
    Arguments:
        folder {str} -- the folder to look at
        file_name {str} -- the file name
        ext {str} -- the file extension
    Returns:
        str -- the final file path

    """
    return os.path.join(folder, f'{file_name}.{ext}')

def tablepop(docx_loc, docx_file, 
             docx_start_row, docx_start_col,
             xlsx_loc, xlsx_list, 
             xlsx_start_row, xlsx_start_col, 
             filled_shells_loc, filled_shells_file, enforce_order):
    
    """Populates a table in word by modifying the file's XML
    
    Arguments:

    Returns:

    """
    
    table_shell = Document(file_loc(docx_loc, docx_file, 'docx'))

    xlsx_list = [xlsxfile for xlsxfile in xlsx_list.split(",")]
    
    print("XLSX_LIST: ", xlsx_list)

    #Loop over the tables of the word document
    for table_index, table in enumerate(table_shell.tables):

        xlsx_file = xlsx_list[table_index]
        
        print("\n"*2, "Populating table", table_index, " with excel file", xlsx_file)

        ##[n] refers to the nth table in the word doc
        doc_tbl = table_shell.tables[table_index]
                
        xltable = pd.read_excel(file_loc(xlsx_loc, xlsx_file, 'xlsx'))
        
        table_adj = 0
        
        rows_to_iterate = xltable.shape[0]
        cols_to_iterate = xltable.shape[1]
        
        print("This table has", cols_to_iterate, "columns", "\n"*2)
        
        #row corresponds to row in the word table, r corresponds to row in the excel file
        for row, r in enumerate(range(xlsx_start_row, rows_to_iterate), start = docx_start_row):
            
            tbl_row = doc_tbl.cell(row + table_adj, 1)

            #if a merged row is found, skip it
            #update while statement -- replace _tc method with non-_ method
            while tbl_row._tc.left == 0:
                table_adj += 1
                print("Table adjustment incremented by 1: ", table_adj)
                tbl_row = doc_tbl.cell(row + table_adj,1)
                
            print("Row:", row, end = " ")
            
            #col corresonds to column in the word table, c corresopnds to column in the excel file
            for col, c in enumerate(range(xlsx_start_col, cols_to_iterate), start = docx_start_col):        
                
                text = str(xltable.iloc[r,c])
                if text.upper() == "NAN":
                    text = ""

                paragraph = doc_tbl.cell(row + table_adj, col).paragraphs[0]

                #Clear out cell contents before populating
                paragraph.clear()
                
                #Using runs preserves table shell format                
                run = paragraph.add_run()
                
                run.text = text

                #update this so 5 is a variable
                print(" ", text.rjust(5, ' '), " ", end = " ")

            print()
                
        table_shell.save(file_loc(filled_shells_loc, filled_shells_file, 'docx'))
        
    return print("Finished populating: ", filled_shells_file)    


if __name__ == '__main__':
    
    parser = argparse.ArgumentParser(description = 'Use excel output to populate table shells in word')

    parser.add_argument('docx_loc', type = str, help = 'Word table shell location')
    parser.add_argument('docx_file', type = str,  help = 'Word table shell filename')
    parser.add_argument('docx_start_row', type = int, help = "The row containing the first cell to populate in the word table shells")
    parser.add_argument('docx_start_col', type = int, help = "The column containing the first cell to populate in the word table shells")        
    parser.add_argument('xlsx_loc', type = str, help = 'Excel file(s) location')
    parser.add_argument('xlsx_list', type = str, help = 'Comma separated list of excel files (no white space between files)')
    parser.add_argument('xlsx_start_row', type = int, help = "The row containing the first cell to pull from the xlsx table")
    parser.add_argument('xlsx_start_col', type = int, help = "The column containing the first cell to pull from the xlsx table")
    parser.add_argument('filled_shells_loc', type = str, help = 'Filled shells save location')
    parser.add_argument('filled_shells_file', type = str, help = 'Filled shells save filename')
    parser.add_argument('enforce_order', type = str, help = 'Enforce order betweeen table shell and excel output via descriptors')
    
    args = parser.parse_args()

    tablepop(args.docx_loc, args.docx_file, 
             args.docx_start_row, args.docx_start_col, 
             args.xlsx_loc, args.xlsx_list, 
             args.xlsx_start_row, args.xlsx_start_col,
             args.filled_shells_loc, args.filled_shells_file, args.enforce_order)
    
